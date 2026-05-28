"""Tests for convert_ops endpoints with full header + file-open verification.

docxedit: fast, no external services.
pdf2docx / docx2pdf: @pytest.mark.integration, skip when LibreOffice absent.
"""
import io
import shutil

import pytest

from conftest import (
    AUTH_HEADERS,
    assert_valid_docx,
    assert_content_type,
    assert_content_disposition,
    assert_x_header,
)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _make_docx(text: str = "Hello world") -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ─── /api/docxedit ────────────────────────────────────────────────────────────

class TestDocxEdit:
    def test_replace_word_returns_docx_with_x_info(self, client):
        pytest.importorskip("docx", reason="python-docx not installed")
        docx_bytes = _make_docx("Hello world, goodbye world.")
        r = client.post(
            "/api/docxedit",
            files=[("file", ("doc.docx", docx_bytes, _DOCX_CT))],
            data={"find": "world", "replace": "EduBot"},
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 200
        assert_content_type(r.headers, "word")
        assert_content_disposition(r.headers, ".docx")
        assert_x_header(r.headers, "x-info")
        assert_valid_docx(r.content)

    def test_replace_count_reflected_in_x_info(self, client):
        pytest.importorskip("docx", reason="python-docx not installed")
        docx_bytes = _make_docx("world world world — three occurrences.")
        r = client.post(
            "/api/docxedit",
            files=[("file", ("doc.docx", docx_bytes, _DOCX_CT))],
            data={"find": "world", "replace": "EduBot"},
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 200
        info = assert_x_header(r.headers, "x-info")
        # Should mention the count of replacements
        assert any(c.isdigit() for c in info), f"Expected a number in x-info: {info!r}"

    def test_replacement_actually_applied(self, client):
        pytest.importorskip("docx", reason="python-docx not installed")
        from docx import Document
        docx_bytes = _make_docx("The quick brown fox.")
        r = client.post(
            "/api/docxedit",
            files=[("file", ("doc.docx", docx_bytes, _DOCX_CT))],
            data={"find": "fox", "replace": "cat"},
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 200
        # Re-open the returned DOCX and verify the replacement
        doc = Document(io.BytesIO(r.content))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "cat" in full_text
        assert "fox" not in full_text

    def test_not_docx_returns_error(self, client):
        r = client.post(
            "/api/docxedit",
            files=[("file", ("img.png", b"\x89PNG\r\n\x1a\n", "image/png"))],
            data={"find": "hello", "replace": "bye"},
            headers=AUTH_HEADERS,
        )
        assert r.status_code in (400, 415, 422, 500)
        assert r.status_code != 200

    def test_too_large_returns_413(self, client, too_large_bytes):
        r = client.post(
            "/api/docxedit",
            files=[("file", ("big.docx", too_large_bytes, _DOCX_CT))],
            data={"find": "x", "replace": "y"},
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 413

    def test_empty_find_returns_error(self, client):
        pytest.importorskip("docx", reason="python-docx not installed")
        docx_bytes = _make_docx("Some text here.")
        r = client.post(
            "/api/docxedit",
            files=[("file", ("doc.docx", docx_bytes, _DOCX_CT))],
            data={"find": "", "replace": "something"},
            headers=AUTH_HEADERS,
        )
        assert r.status_code in (400, 422)


# ─── /api/pdf2docx ────────────────────────────────────────────────────────────

@pytest.mark.integration
class TestPdf2Docx:
    @pytest.fixture(autouse=True)
    def require_libreoffice(self):
        if not (shutil.which("libreoffice") or shutil.which("soffice")):
            pytest.skip("LibreOffice not available")

    def test_pdf2docx_returns_docx_with_headers(self, client, table_pdf):
        r = client.post(
            "/api/pdf2docx",
            files=[("file", ("table.pdf", table_pdf, "application/pdf"))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 200
        assert_content_type(r.headers, "word")
        assert_content_disposition(r.headers, "converted.docx")
        assert_x_header(r.headers, "x-info")
        assert_x_header(r.headers, "x-page-count")
        assert_valid_docx(r.content)

    def test_pdf2docx_cyrillic_pdf_returns_docx(self, client, cyrillic_pdf):
        r = client.post(
            "/api/pdf2docx",
            files=[("file", ("cyrillic.pdf", cyrillic_pdf, "application/pdf"))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 200
        assert_valid_docx(r.content)

    def test_pdf2docx_not_pdf_returns_error(self, client):
        r = client.post(
            "/api/pdf2docx",
            files=[("file", ("img.png", b"\x89PNG\r\n\x1a\n", "image/png"))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code in (400, 415, 422, 500)

    def test_pdf2docx_password_pdf_returns_error(self, client, password_pdf):
        r = client.post(
            "/api/pdf2docx",
            files=[("file", ("locked.pdf", password_pdf, "application/pdf"))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code in (400, 422, 500)

    def test_pdf2docx_too_large_returns_413(self, client, too_large_bytes):
        r = client.post(
            "/api/pdf2docx",
            files=[("file", ("big.pdf", too_large_bytes, "application/pdf"))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 413


# ─── /api/docx2pdf ────────────────────────────────────────────────────────────

@pytest.mark.integration
class TestDocx2Pdf:
    @pytest.fixture(autouse=True)
    def require_libreoffice(self):
        if not (shutil.which("libreoffice") or shutil.which("soffice")):
            pytest.skip("LibreOffice not available")

    def test_docx2pdf_returns_pdf_with_info(self, client):
        pytest.importorskip("docx", reason="python-docx not installed")
        docx_bytes = _make_docx("Test document for conversion to PDF.")
        r = client.post(
            "/api/docx2pdf",
            files=[("file", ("doc.docx", docx_bytes, _DOCX_CT))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 200
        assert_content_type(r.headers, "pdf")
        assert_x_header(r.headers, "x-info")
        assert r.content[:4] == b"%PDF"

    def test_docx2pdf_not_docx_returns_error(self, client):
        r = client.post(
            "/api/docx2pdf",
            files=[("file", ("img.png", b"\x89PNG\r\n\x1a\n", "image/png"))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code in (400, 415, 422, 500)

    def test_docx2pdf_too_large_returns_413(self, client, too_large_bytes):
        r = client.post(
            "/api/docx2pdf",
            files=[("file", ("big.docx", too_large_bytes, _DOCX_CT))],
            headers=AUTH_HEADERS,
        )
        assert r.status_code == 413
