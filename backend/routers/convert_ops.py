# EduBot — Document conversion routes
# pdf2docx (multi-engine: libreoffice → pdf2docx → pymupdf;
#           marker/docling auto-enabled if their packages are installed)
# docx2pdf (libreoffice → mammoth+weasyprint → reportlab)
# docxedit  (find-and-replace in .docx)

import io
import os
import re
import asyncio
import functools
import zipfile
import threading
import tempfile
import time

from fastapi import APIRouter, File, UploadFile, Form, Request, HTTPException
from fastapi.responses import Response
from loguru import logger
from typing import Optional, List
from urllib.parse import quote as url_quote

from shared import (
    limiter, _io_pool,
    safe_header, is_image_bytes,
    read_upload, read_uploads,
    safe_url_fetcher,
    FONT_REGULAR, FONT_BOLD,
    _get_user_id,
)

router = APIRouter()


def _run_soffice(args: list, env: dict, timeout: int = 60):
    """Run a headless LibreOffice command in its OWN process group.

    `soffice` forks detached `soffice.bin`/`oosplash` children; a plain
    subprocess timeout kills only the parent, leaving orphans that hold the
    temp profile open and leak RAM. Starting a new session lets us SIGKILL the
    whole group on timeout. Raises subprocess.TimeoutExpired on timeout.
    """
    import subprocess, signal
    proc = subprocess.Popen(
        args, env=env,
        stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        start_new_session=True,
    )
    try:
        out, err = proc.communicate(timeout=timeout)
        return proc.returncode, out, err
    except subprocess.TimeoutExpired:
        try:
            os.killpg(os.getpgid(proc.pid), signal.SIGKILL)
        except Exception:
            proc.kill()
        try:
            proc.communicate(timeout=5)
        except Exception:
            pass
        raise


# ─── PDF → DOCX: helpers (ml-chain: marker → docling → libreoffice → pdf2docx → pymupdf) ──

def _do_pdf2docx(pdf_path: str, docx_path: str) -> None:
    from pdf2docx import Converter
    cv = Converter(pdf_path)
    try:
        cv.convert(docx_path, multi_processing=False, start=0)
    finally:
        cv.close()


_marker_converter = None
_marker_lock = threading.Lock()

def _get_marker_converter():
    global _marker_converter
    if _marker_converter is None:
        with _marker_lock:
            if _marker_converter is None:
                from marker.converters.pdf import PdfConverter
                from marker.models import create_model_dict
                _marker_converter = PdfConverter(artifact_dict=create_model_dict())
                logger.info("marker-pdf model loaded")
    return _marker_converter


def _do_pdf2docx_marker(pdf_path: str, docx_path: str) -> None:
    from marker.output import text_from_rendered
    import pypandoc
    converter = _get_marker_converter()
    rendered = converter(pdf_path)
    md_text, _, _ = text_from_rendered(rendered)
    if not md_text or len(md_text.strip()) < 10:
        raise RuntimeError("marker bo'sh natija qaytardi")
    pypandoc.convert_text(
        md_text, to="docx", format="md", outputfile=docx_path,
        extra_args=["--standalone"],
    )


def _do_pdf2docx_libreoffice(pdf_path: str, docx_path: str) -> None:
    import shutil
    if shutil.which("libreoffice") is None and shutil.which("soffice") is None:
        raise RuntimeError("LibreOffice o'rnatilmagan")
    outdir = os.path.dirname(docx_path)
    lo_home = tempfile.mkdtemp(prefix="lo_")
    env = os.environ.copy()
    env["HOME"] = lo_home
    # Serialize LibreOffice — otherwise up to _IO_WORKERS soffice instances run
    # at once (~150-300MB each) and OOM the basic tier.
    with _office_convert_sem:
        try:
            rc, _out, _err = _run_soffice(
                ["libreoffice", "--headless", "--norestore",
                 "--convert-to", "docx", "--infilter=writer_pdf_import",
                 "--outdir", outdir, pdf_path],
                env=env, timeout=60,
            )
        finally:
            shutil.rmtree(lo_home, ignore_errors=True)
    if rc != 0:
        stderr = (_err or b"").decode(errors="replace")[:300]
        raise RuntimeError(f"libreoffice exit {rc}: {stderr}")
    expected = os.path.join(outdir, os.path.splitext(os.path.basename(pdf_path))[0] + ".docx")
    if not os.path.exists(expected) or os.path.getsize(expected) < 100:
        raise RuntimeError("LibreOffice DOCX chiqarilmadi yoki bo'sh")
    if os.path.abspath(expected) != os.path.abspath(docx_path):
        shutil.move(expected, docx_path)


def _do_pdf2docx_pymupdf(pdf_path: str, docx_path: str) -> None:
    import fitz
    from docx import Document
    from docx.shared import Pt, RGBColor
    pdf = fitz.open(pdf_path)
    try:
        _pymupdf_extract(pdf, docx_path)
    finally:
        pdf.close()


def _pymupdf_extract(pdf, docx_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc_out = Document()

    sample_sizes = []
    for page in pdf[:3]:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        sz = span.get("size", 0)
                        if sz > 0 and span.get("text", "").strip():
                            sample_sizes.append(sz)
                            if len(sample_sizes) >= 500:
                                break
                    if len(sample_sizes) >= 500:
                        break
            if len(sample_sizes) >= 500:
                break
        if len(sample_sizes) >= 500:
            break
    body_size = sorted(sample_sizes)[len(sample_sizes) // 2] if sample_sizes else 11.0

    for page_num, page in enumerate(pdf):
        if page_num > 0:
            doc_out.add_paragraph()

        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            all_spans = [
                s for ln in block.get("lines", [])
                for s in ln.get("spans", [])
                if s.get("text", "").strip()
            ]
            if not all_spans:
                continue

            avg_size = sum(s["size"] for s in all_spans) / len(all_spans)
            block_text = " ".join(s["text"].strip() for s in all_spans)

            if avg_size >= body_size * 1.6:
                doc_out.add_heading(block_text, level=1)
            elif avg_size >= body_size * 1.35:
                doc_out.add_heading(block_text, level=2)
            elif avg_size >= body_size * 1.15:
                doc_out.add_heading(block_text, level=3)
            else:
                p = doc_out.add_paragraph()
                for span in all_spans:
                    text = span.get("text", "")
                    if not text:
                        continue
                    run = p.add_run(text)
                    flags = span.get("flags", 0)
                    run.bold = bool(flags & 16)
                    run.italic = bool(flags & 2)
                    sz = span.get("size", body_size)
                    if 4 <= sz <= 72:
                        run.font.size = Pt(round(sz))
                    color = span.get("color", 0)
                    if color:
                        r, g, b = (color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF
                        if (r, g, b) != (0, 0, 0):
                            try:
                                run.font.color.rgb = RGBColor(r, g, b)
                            except Exception:
                                pass

    doc_out.save(docx_path)


def _do_pdf2docx_ocr(pdf_path: str, docx_path: str) -> None:
    import fitz
    import pytesseract
    from PIL import Image as PILImage
    from docx import Document

    pdf = fitz.open(pdf_path)
    try:
        doc_out = Document()
        doc_out.add_heading("OCR natijasi", level=1)
        for page_num, page in enumerate(pdf):
            if page_num > 0:
                doc_out.add_page_break()
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
            try:
                try:
                    text = pytesseract.image_to_string(img, lang="uzb+rus+eng", timeout=30)
                except Exception:
                    text = pytesseract.image_to_string(img, lang="eng", timeout=30)
            finally:
                img.close()
            for line in (text or "").split("\n"):
                line = line.strip()
                if line:
                    doc_out.add_paragraph(line)
        doc_out.save(docx_path)
    finally:
        pdf.close()


def _do_pdf2docx_docling(pdf_path: str, docx_path: str) -> None:
    from docling.document_converter import DocumentConverter
    converter = DocumentConverter()
    result = converter.convert(pdf_path)
    doc = result.document
    saved = False
    try:
        docx_bytes = doc.export_to_docx()
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        saved = True
    except AttributeError:
        pass
    if not saved:
        try:
            from docling.backend.docx_backend import DocxDocumentBackend
            with open(docx_path, "wb") as f:
                DocxDocumentBackend().write(doc, f)
            saved = True
        except Exception:
            pass
    if not saved:
        raise RuntimeError("docling DOCX saqlash usuli topilmadi")


def _do_pdf2docx_best(pdf_path: str, docx_path: str, is_scanned: bool) -> str:
    def _disabled(name: str) -> bool:
        # marker/docling load multi-GB ML models → instant OOM on Railway's
        # basic tier. They are OFF by default and only used when BOTH the
        # explicit opt-in env (MARKER_ENABLED=1 / DOCLING_ENABLED=1) is set
        # AND the package is importable.
        if os.environ.get(f"{name}_DISABLED", "").lower() in ("1", "true", "yes"):
            return True
        pkg = {"MARKER": "marker", "DOCLING": "docling"}.get(name)
        if pkg:
            if os.environ.get(f"{name}_ENABLED", "").lower() not in ("1", "true", "yes"):
                return True  # default off — opt-in required
            try:
                __import__(pkg)
            except ImportError:
                return True
        return False

    try:
        import fitz
        with fitz.open(pdf_path) as _pdf:
            page_count = _pdf.page_count
    except Exception:
        page_count = 999

    try:
        MARKER_MAX_PAGES = int(os.environ.get("MARKER_MAX_PAGES", "10"))
    except ValueError:
        MARKER_MAX_PAGES = 10

    if is_scanned:
        chain = [
            ("ocr-tesseract", lambda: _do_pdf2docx_ocr(pdf_path, docx_path)),
            ("pdf2docx",      lambda: _do_pdf2docx(pdf_path, docx_path)),
        ]
    else:
        # Matn-PDF lar uchun tartib o'zgartirildi:
        # pdf2docx paketi ro'yxatlar, jadvallar va matn formatlashni eng to'g'ri
        # saqlaydi (foydalanuvchi xabar bergan "ro'yxatlar nuqta va bo'sh joylar bilan"
        # muammosi LibreOffice writer_pdf_import filtri sababli edi). Endi pdf2docx
        # birinchi sinaladi, LibreOffice ikkinchi (umumiy fallback), pymupdf oxirgi.
        chain = []
        if not _disabled("MARKER") and page_count <= MARKER_MAX_PAGES:
            chain.append(("marker", lambda: _do_pdf2docx_marker(pdf_path, docx_path)))
        if not _disabled("DOCLING"):
            chain.append(("docling", lambda: _do_pdf2docx_docling(pdf_path, docx_path)))
        chain.extend([
            ("pdf2docx",      lambda: _do_pdf2docx(pdf_path, docx_path)),
            ("libreoffice",   lambda: _do_pdf2docx_libreoffice(pdf_path, docx_path)),
            ("pymupdf",       lambda: _do_pdf2docx_pymupdf(pdf_path, docx_path)),
        ])

    last_err = None
    for name, fn in chain:
        try:
            fn()
            if os.path.exists(docx_path) and os.path.getsize(docx_path) >= 2000:
                return name
        except Exception as e:
            last_err = e
            logger.warning(f"pdf2docx [{name}] muvaffaqiyatsiz: {type(e).__name__}: {str(e)[:120]}")
        try:
            if os.path.exists(docx_path):
                os.unlink(docx_path)
        except Exception:
            pass

    raise RuntimeError(
        f"Barcha usullar muvaffaqiyatsiz: "
        f"{type(last_err).__name__ if last_err else 'Unknown'}: {str(last_err)[:100]}"
    )


@router.post("/api/pdf2docx")
@limiter.limit("10/minute")
async def pdf_to_docx(request: Request, file: UploadFile = File(...)):
    t0 = time.time()
    user_id = _get_user_id(request)
    # Use OS temp dir (cross-platform; /tmp is Linux-only) with random subdir
    # so concurrent requests can't collide and one finally cleans up everything.
    tmpdir    = tempfile.mkdtemp(prefix="edubot_pdf2docx_")
    pdf_path  = os.path.join(tmpdir, "input.pdf")
    docx_path = os.path.join(tmpdir, "output.docx")
    try:
        data = await read_upload(file, "/api/pdf2docx")

        from pypdf import PdfReader
        try:
            reader = PdfReader(io.BytesIO(data))
            page_count = len(reader.pages)
        except Exception:
            raise HTTPException(status_code=422,
                detail="PDF fayl ochib bo'lmadi. Fayl buzilgan yoki parol bilan himoyalangan.")

        if page_count == 0:
            raise HTTPException(status_code=422, detail="PDF bo'sh (0 sahifa).")

        MAX_PAGES = 50
        if page_count > MAX_PAGES:
            raise HTTPException(status_code=400,
                detail=f"PDF {page_count} sahifali. Maksimal {MAX_PAGES} sahifa. "
                       f"PDF'ni bo'lib (split) qayta urinib ko'ring.")

        sample = min(3, page_count)
        total_chars = sum(
            len((reader.pages[i].extract_text() or "").strip())
            for i in range(sample)
        )
        is_scanned = total_chars < 20

        with open(pdf_path, "wb") as f:
            f.write(data)

        loop = asyncio.get_running_loop()
        try:
            conv_method = await asyncio.wait_for(
                loop.run_in_executor(_io_pool, _do_pdf2docx_best, pdf_path, docx_path, is_scanned),
                timeout=120.0,
            )
        except asyncio.TimeoutError:
            raise HTTPException(status_code=408,
                detail=f"Konversiya {page_count} sahifa uchun 120 soniyadan oshdi.")

        if not os.path.exists(docx_path):
            raise HTTPException(status_code=500, detail="Konversiya muvaffaqiyatsiz (fayl yaratilmadi).")

        with open(docx_path, "rb") as f:
            out = f.read()

        if len(out) < 2000:
            raise HTTPException(status_code=422,
                detail="Konversiya natijasi bo'sh. PDF tuzilishi qo'llab-quvvatlanmaydi.")

        info_parts = [f"{page_count} sahifa aylantirildi ({conv_method})"]
        if is_scanned:
            info_parts.append("OCR orqali")
        info = " - ".join(info_parts)

        elapsed = time.time() - t0
        logger.info(f"pdf2docx [{conv_method}]: {page_count}p {'scan' if is_scanned else 'text'}, "
                    f"{len(data)//1024}KB → {len(out)//1024}KB, {elapsed:.1f}s")

        return Response(
            content=out,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=converted.docx",
                "X-Info": safe_header(info),
                "X-Page-Count": str(page_count),
            },
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"pdf2docx xato: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500,
            detail="Xizmatda kutilmagan xatolik yuz berdi. Iltimos, qayta urinib ko'ring.")
    finally:
        # Wipe the whole tempdir (handles partial writes / subprocess artifacts too)
        import shutil as _shutil
        _shutil.rmtree(tmpdir, ignore_errors=True)

# ─── DOCX → PDF: helpers ──────────────────────────────────────────────────────

def _rl_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    try:
        pdfmetrics.registerFont(TTFont('DV', FONT_REGULAR))
        pdfmetrics.registerFont(TTFont('DV-Bold', FONT_BOLD))
        return 'DV', 'DV-Bold'
    except Exception:
        return 'Helvetica', 'Helvetica-Bold'


_WORD_TO_PDF_EXTS = {".doc", ".docx"}


def _env_int(name: str, default: int, minimum: int, maximum: int) -> int:
    try:
        value = int(os.environ.get(name, str(default)))
    except Exception:
        value = default
    return max(minimum, min(maximum, value))


_OFFICE_CONVERT_TIMEOUT = _env_int("OFFICE_CONVERT_TIMEOUT", 120, 30, 300)
_OFFICE_CONVERT_SLOTS   = _env_int("OFFICE_CONVERT_SLOTS", 2, 1, 4)
_office_convert_sem     = threading.BoundedSemaphore(_OFFICE_CONVERT_SLOTS)


def _looks_like_docx(data: bytes) -> bool:
    if not data.startswith(b"PK"):
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = set(zf.namelist())
        return "[Content_Types].xml" in names and any(n.startswith("word/") for n in names)
    except Exception:
        return False


def _looks_like_legacy_doc(data: bytes) -> bool:
    sample = data[:512].lstrip().lower()
    return (
        data.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
        or sample.startswith(b"{\\rtf")
        or sample.startswith(b"<html")
        or b"<html" in sample[:128]
    )


def _detect_word_extension(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext not in _WORD_TO_PDF_EXTS:
        if _looks_like_docx(data):
            return ".docx"
        if _looks_like_legacy_doc(data):
            return ".doc"
        raise HTTPException(status_code=415, detail="Faqat .doc yoki .docx fayl qabul qilinadi.")
    if ext == ".docx":
        if _looks_like_docx(data):
            return ".docx"
        if _looks_like_legacy_doc(data):
            return ".doc"
        raise HTTPException(status_code=422, detail="DOCX fayl buzilgan yoki noto'g'ri formatda.")
    if _looks_like_docx(data):
        return ".docx"
    if not _looks_like_legacy_doc(data):
        raise HTTPException(status_code=422, detail="DOC fayl buzilgan yoki noto'g'ri formatda.")
    return ".doc"


def _safe_pdf_filename(filename: str) -> str:
    stem = os.path.splitext(os.path.basename(filename or ""))[0]
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._-")
    stem = stem[:64] or "document"
    return f"{stem}.pdf"


def _pdf_content_disposition(filename: str) -> str:
    """attachment header with ASCII fallback + RFC 5987 filename* so a
    Cyrillic/Uzbek original name survives instead of collapsing to document.pdf."""
    ascii_name = _safe_pdf_filename(filename)
    stem = os.path.splitext(os.path.basename(filename or ""))[0][:64].strip() or "document"
    utf8_name = url_quote(f"{stem}.pdf")
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{utf8_name}"


def _inspect_docx(data: bytes) -> dict:
    if not _looks_like_docx(data):
        return {"readable": False, "paragraphs": None, "tables": None, "images": None}
    try:
        from docx import Document as _DocCheck
        doc = _DocCheck(io.BytesIO(data))
        para_count = sum(1 for p in doc.paragraphs if p.text.strip())
        table_count = len(doc.tables)
        image_count = sum(
            1 for rel in doc.part.rels.values()
            if "image" in getattr(rel, "reltype", "")
        )
        return {"readable": True, "paragraphs": para_count, "tables": table_count, "images": image_count}
    except Exception:
        return {"readable": False, "paragraphs": None, "tables": None, "images": None}


def _validate_pdf_bytes(pdf_bytes: bytes) -> tuple:
    if not pdf_bytes or len(pdf_bytes) < 500:
        raise RuntimeError("PDF natija bo'sh yoki juda kichik")
    if not pdf_bytes[:1024].lstrip().startswith(b"%PDF"):
        raise RuntimeError("PDF natija noto'g'ri formatda")
    import fitz
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        page_count = pdf.page_count
        if page_count < 1:
            raise RuntimeError("PDF sahifa topilmadi")
        text_chars = 0
        for i in range(min(page_count, 3)):
            text_chars += len((pdf[i].get_text("text") or "").strip())
        return page_count, text_chars
    finally:
        pdf.close()


def _read_lo_pdf_output(outdir: str, input_stem: str) -> bytes:
    expected = os.path.join(outdir, input_stem + ".pdf")
    candidates = [expected]
    try:
        candidates.extend(
            os.path.join(outdir, name)
            for name in os.listdir(outdir)
            if name.lower().endswith(".pdf") and os.path.join(outdir, name) != expected
        )
    except Exception:
        pass
    for path in candidates:
        if os.path.exists(path) and os.path.getsize(path) > 500:
            with open(path, "rb") as f:
                return f.read()
    raise RuntimeError("LibreOffice PDF chiqarilmadi yoki bo'sh")


def _do_docx2pdf(data: bytes, fn_regular: str, fn_bold: str) -> bytes:
    """
    CPU-bound DOCX→PDF via python-docx + ReportLab.
    Preserves headings, bold/italic/underline/color/size, tables, images, lists.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, PageBreak,
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from PIL import Image as PILImage

    doc = Document(io.BytesIO(data))
    PAGE_W = A4[0] - 4 * cm

    ALIGN_MAP = {
        WD_ALIGN_PARAGRAPH.LEFT:    TA_LEFT,
        WD_ALIGN_PARAGRAPH.CENTER:  TA_CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT:   TA_RIGHT,
        WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY,
        None: TA_LEFT,
    }

    S = {
        'h1':  ParagraphStyle('h1',  fontName=fn_bold,    fontSize=20, leading=26, spaceAfter=8,  spaceBefore=18),
        'h2':  ParagraphStyle('h2',  fontName=fn_bold,    fontSize=16, leading=21, spaceAfter=6,  spaceBefore=14),
        'h3':  ParagraphStyle('h3',  fontName=fn_bold,    fontSize=13, leading=17, spaceAfter=5,  spaceBefore=10),
        'h4':  ParagraphStyle('h4',  fontName=fn_bold,    fontSize=12, leading=15, spaceAfter=4,  spaceBefore=8),
        'body':ParagraphStyle('body',fontName=fn_regular, fontSize=11, leading=16, spaceAfter=6),
        'lb':  ParagraphStyle('lb',  fontName=fn_regular, fontSize=11, leading=16, spaceAfter=3,  leftIndent=18),
        'ln':  ParagraphStyle('ln',  fontName=fn_regular, fontSize=11, leading=16, spaceAfter=3,  leftIndent=18),
        'tc':  ParagraphStyle('tc',  fontName=fn_regular, fontSize=9,  leading=13),
        'tch': ParagraphStyle('tch', fontName=fn_bold,    fontSize=9,  leading=13),
    }

    _style_ctr = [0]

    def derived(base, alignment, extra_indent=0):
        _style_ctr[0] += 1
        return ParagraphStyle(
            f'_d{_style_ctr[0]}', parent=base,
            alignment=alignment, leftIndent=base.leftIndent + extra_indent,
        )

    def esc(t: str) -> str:
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def runs_to_markup(para) -> str:
        if not para.runs:
            return esc(para.text or "")
        parts = []
        for run in para.runs:
            if not run.text:
                continue
            t = esc(run.text)
            try:
                clr = run.font.color
                if clr and clr.type and clr.rgb:
                    t = f'<font color="#{clr.rgb}">{t}</font>'
            except Exception:
                pass
            try:
                if run.font.size and run.font.size.pt:
                    fs = round(run.font.size.pt)
                    if fs > 0 and abs(fs - 11) > 1:
                        t = f'<font size="{fs}">{t}</font>'
            except Exception:
                pass
            if run.bold:      t = f'<b>{t}</b>'
            if run.italic:    t = f'<i>{t}</i>'
            if run.underline: t = f'<u>{t}</u>'
            parts.append(t)
        result = "".join(parts)
        return result if result.strip() else esc(para.text or "")

    def heading_level(element, para) -> int:
        style_name = (para.style.name or "").lower()
        for n in (1, 2, 3, 4, 5, 6):
            if f"heading {n}" in style_name or f"заголовок {n}" in style_name:
                return min(n, 4)
        try:
            pPr = element.find(qn("w:pPr"))
            if pPr is not None:
                ol = pPr.find(qn("w:outlineLvl"))
                if ol is not None:
                    lvl = int(ol.get(qn("w:val"), 9))
                    if lvl <= 3:
                        return lvl + 1
        except Exception:
            pass
        return 0

    images: dict = {}
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                images[rel.rId] = rel.target_part.blob
    except Exception:
        pass

    list_counters: dict = {}
    story = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            try:
                para = DocxParagraph(element, doc)
                try:
                    style_name = (para.style.name or "").lower()
                except Exception:
                    style_name = ""

                for blip in element.findall(".//" + qn("a:blip")):
                    rId = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
                    if rId and rId in images:
                        try:
                            with PILImage.open(io.BytesIO(images[rId])) as pil:
                                iw, ih = pil.size
                                if iw > PAGE_W:
                                    ratio = PAGE_W / iw
                                    iw, ih = PAGE_W, ih * ratio
                                out_buf = io.BytesIO()
                                pil.convert("RGB").save(out_buf, format="PNG")
                            out_buf.seek(0)
                            story.append(RLImage(out_buf, width=iw, height=ih))
                            story.append(Spacer(1, 6))
                        except Exception as e:
                            logger.debug(f"docx image render skip: {e}")

                markup = runs_to_markup(para)
                if not markup.strip():
                    story.append(Spacer(1, 3))
                    continue

                try:
                    align = ALIGN_MAP.get(para.alignment, TA_LEFT)
                except Exception:
                    align = TA_LEFT
                hlvl = heading_level(element, para)

                if hlvl == 1:
                    st = derived(S["h1"], align)
                elif hlvl == 2:
                    st = derived(S["h2"], align)
                elif hlvl == 3:
                    st = derived(S["h3"], align)
                elif hlvl >= 4:
                    st = derived(S["h4"], align)
                elif "list bullet" in style_name or "list paragraph" in style_name:
                    markup = f"• {markup}"
                    st = derived(S["lb"], align)
                elif "list number" in style_name:
                    try:
                        numId = element.find(".//" + qn("w:numId"))
                        key = numId.get(qn("w:val"), "0") if numId is not None else "0"
                    except Exception:
                        key = "0"
                    list_counters[key] = list_counters.get(key, 0) + 1
                    markup = f"{list_counters[key]}. {markup}"
                    st = derived(S["ln"], align)
                else:
                    st = derived(S["body"], align)

                try:
                    story.append(Paragraph(markup, st))
                except Exception:
                    try:
                        story.append(Paragraph(esc(para.text or ""), S["body"]))
                    except Exception:
                        pass
            except Exception:
                pass

        elif tag == "tbl":
            try:
                table = DocxTable(element, doc)
                rows_data = []
                for r_idx, row in enumerate(table.rows):
                    try:
                        row_cells = []
                        is_header = r_idx == 0
                        for cell in row.cells:
                            try:
                                cell_para_text = " ".join(p.text for p in cell.paragraphs).strip()
                            except Exception:
                                cell_para_text = ""
                            row_cells.append(Paragraph(esc(cell_para_text),
                                                       S["tch"] if is_header else S["tc"]))
                        rows_data.append(row_cells)
                    except Exception:
                        pass

                if rows_data:
                    col_n = max(len(r) for r in rows_data)
                    col_w = PAGE_W / col_n if col_n else PAGE_W
                    tbl = Table(rows_data, colWidths=[col_w] * col_n, repeatRows=1)
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND",    (0, 0), (-1,  0), colors.HexColor("#e8e8e8")),
                        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f7f7")]),
                        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#bbbbbb")),
                        ("TOPPADDING",    (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
                        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                        ("WORDWRAP",      (0, 0), (-1, -1), "WORD"),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 10))
            except Exception:
                pass

        elif tag == "sectPr":
            pass

    if not story:
        story.append(Paragraph("(Bo'sh hujjat)", S["body"]))

    out = io.BytesIO()
    pdf_doc = SimpleDocTemplate(
        out, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=doc.core_properties.title or "EduBot Document",
        author="EduBot",
    )
    try:
        pdf_doc.build(story)
    except Exception:
        safe = [item for item in story if isinstance(item, (Paragraph, Spacer, PageBreak))]
        if not safe:
            safe = [Paragraph("(Hujjatni aylantirish qisman bajarildi)", S["body"])]
        out2 = io.BytesIO()
        pdf_doc2 = SimpleDocTemplate(
            out2, pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm,
            topMargin=2.5 * cm, bottomMargin=2.5 * cm,
            title="EduBot Document", author="EduBot",
        )
        pdf_doc2.build(safe)
        return out2.getvalue()
    return out.getvalue()


def _do_docx2pdf_libreoffice(data: bytes, ext: str) -> bytes:
    import subprocess, shutil
    from pathlib import Path

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice topilmadi")

    ext = ext if ext in _WORD_TO_PDF_EXTS else ".docx"
    with _office_convert_sem:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_stem = "input"
            input_path = os.path.join(tmpdir, input_stem + ext)
            with open(input_path, "wb") as f:
                f.write(data)

            profile_dir = tempfile.mkdtemp(prefix="lo_profile_")
            try:
                env = os.environ.copy()
                env["HOME"] = profile_dir
                env.setdefault("SAL_USE_VCLPLUGIN", "svp")
                env.setdefault("LC_ALL", "C.UTF-8")
                profile_uri = Path(profile_dir).resolve().as_uri()

                configured_filter = os.environ.get("OFFICE_PDF_EXPORT_FILTER", "").strip()
                export_filters = []
                for candidate in (configured_filter, "pdf:writer_pdf_Export", "pdf"):
                    if candidate and candidate not in export_filters:
                        export_filters.append(candidate)
                errors = []
                for export_filter in export_filters:
                    for name in os.listdir(tmpdir):
                        if name.lower().endswith(".pdf"):
                            try:
                                os.remove(os.path.join(tmpdir, name))
                            except Exception:
                                pass

                    cmd = [
                        soffice, "--headless", "--invisible", "--nodefault",
                        "--nologo", "--nofirststartwizard", "--norestore",
                        f"-env:UserInstallation={profile_uri}",
                        "--convert-to", export_filter,
                        "--outdir", tmpdir, input_path,
                    ]
                    try:
                        rc, _out, _err = _run_soffice(
                            cmd, env=env, timeout=_OFFICE_CONVERT_TIMEOUT,
                        )
                    except subprocess.TimeoutExpired:
                        raise RuntimeError(f"LibreOffice timeout {_OFFICE_CONVERT_TIMEOUT}s")

                    stderr = (_err or b"").decode(errors="replace")[:300]
                    stdout = (_out or b"").decode(errors="replace")[:300]
                    if rc != 0:
                        errors.append(f"{export_filter}: exit {rc}: {stderr or stdout}")
                        continue

                    try:
                        pdf_bytes = _read_lo_pdf_output(tmpdir, input_stem)
                        _validate_pdf_bytes(pdf_bytes)
                        return pdf_bytes
                    except Exception as e:
                        errors.append(f"{export_filter}: {type(e).__name__}: {str(e)[:120]}")

                raise RuntimeError("; ".join(errors[-3:]) or "LibreOffice konvertatsiya qila olmadi")
            finally:
                shutil.rmtree(profile_dir, ignore_errors=True)


def _do_docx2pdf_mammoth(data: bytes) -> bytes:
    import mammoth
    from weasyprint import HTML
    result = mammoth.convert_to_html(io.BytesIO(data))
    html_body = result.value or ""
    if not html_body.strip():
        raise ValueError("mammoth: bo'sh HTML natija")
    html = (
        "<!DOCTYPE html><html><head>"
        "<meta charset='utf-8'>"
        "<style>body{font-family:DejaVu Sans,sans-serif;font-size:11pt;margin:2cm;}"
        "h1{font-size:18pt}h2{font-size:15pt}h3{font-size:13pt}"
        "table{border-collapse:collapse;width:100%}"
        "td,th{border:1px solid #bbb;padding:4px 6px}"
        "</style></head><body>"
        + html_body
        + "</body></html>"
    )
    # SSRF guard: a crafted DOCX can embed <img src="http://internal-host/">
    # via mammoth → WeasyPrint. safe_url_fetcher rejects everything except data:.
    pdf_bytes = HTML(string=html, url_fetcher=safe_url_fetcher).write_pdf()
    if not pdf_bytes or len(pdf_bytes) < 100:
        raise RuntimeError("weasyprint: bo'sh PDF natija")
    return pdf_bytes


def _do_docx2pdf_best(data: bytes, ext: str, fn_regular: str, fn_bold: str) -> tuple:
    attempts = [("libreoffice", lambda: _do_docx2pdf_libreoffice(data, ext))]
    if ext == ".docx":
        attempts.extend([
            ("mammoth",   lambda: _do_docx2pdf_mammoth(data)),
            ("reportlab", lambda: _do_docx2pdf(data, fn_regular, fn_bold)),
        ])
    last_err = None
    for name, fn in attempts:
        try:
            result = fn()
            page_count, text_chars = _validate_pdf_bytes(result)
            return result, name, page_count, text_chars
        except Exception as e:
            last_err = e
            logger.warning(f"docx2pdf [{name}] muvaffaqiyatsiz: {type(e).__name__}: {str(e)[:120]}")
    raise RuntimeError(
        f"Barcha usullar muvaffaqiyatsiz: {type(last_err).__name__}: {str(last_err)[:100]}"
    )


@router.post("/api/docx2pdf")
@limiter.limit("10/minute")
async def docx_to_pdf(request: Request, file: UploadFile = File(...)):
    t0 = time.time()
    user_id = _get_user_id(request)
    try:
        data = await read_upload(file, "/api/docx2pdf")
        fn, fn_bold = _rl_fonts()

        original_name = file.filename or "document"
        ext = _detect_word_extension(original_name, data)
        doc_stats = _inspect_docx(data) if ext == ".docx" else {
            "readable": False, "paragraphs": None, "tables": None, "images": None,
        }

        if doc_stats["readable"]:
            para_count  = doc_stats["paragraphs"] or 0
            table_count = doc_stats["tables"] or 0
            image_count = doc_stats["images"] or 0
            if para_count == 0 and table_count == 0 and image_count == 0:
                raise HTTPException(status_code=422,
                    detail="Hujjat bo'sh (matn, jadval yoki rasm topilmadi).")
        else:
            para_count = table_count = image_count = None

        loop = asyncio.get_running_loop()
        try:
            pdf_bytes, conv_method, page_count, text_chars = await asyncio.wait_for(
                loop.run_in_executor(_io_pool, _do_docx2pdf_best, data, ext, fn, fn_bold),
                timeout=float(_OFFICE_CONVERT_TIMEOUT + 30),
            )
        except asyncio.TimeoutError:
            raise HTTPException(status_code=408,
                detail=f"Konversiya {_OFFICE_CONVERT_TIMEOUT + 30} soniyadan oshdi.")

        if len(pdf_bytes) < 500 or page_count < 1:
            raise HTTPException(status_code=500, detail="Konversiya natijasi bo'sh. Qayta urinib ko'ring.")

        info_parts = [ext.lstrip(".").upper()]
        if para_count is not None:
            info_parts.append(f"{para_count} paragraf")
        if table_count:
            info_parts.append(f"{table_count} jadval")
        if image_count:
            info_parts.append(f"{image_count} rasm")
        info_parts.append(f"{page_count} sahifa")
        info_parts.append(f"aylantirildi ({conv_method})")
        info = " - ".join(info_parts)

        logger.info(f"docx2pdf [{conv_method}/{ext}]: "
                    f"{para_count if para_count is not None else '?'}p, "
                    f"{page_count} pages, {len(data)//1024}KB -> {len(pdf_bytes)//1024}KB, "
                    f"{time.time()-t0:.1f}s")

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": _pdf_content_disposition(original_name),
                "X-Info": safe_header(info),
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"docx2pdf xato: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500,
            detail="Xizmatda kutilmagan xatolik yuz berdi. Iltimos, qayta urinib ko'ring.")


# ─── DOCX Edit ────────────────────────────────────────────────────────────────

def _docxedit_apply_one(doc, find: str, replace: str, case_sensitive: bool) -> int:
    """Loaded Document'ga BITTA find/replace juftligini qo'llaydi.

    Avval `_do_docxedit` ichida edi, endi multi-pair (bir nechta almashtirish)
    versiyasi shu yordamchini qayta-qayta chaqirishi mumkin bo'lishi uchun
    ajratildi.
    """
    import re as _re

    count = 0

    def _replace_in_para(para):
        nonlocal count
        if not para.runs:
            return

        if case_sensitive:
            hit = 0
            for r in para.runs:
                if find in r.text:
                    hit += r.text.count(find)
                    r.text = r.text.replace(find, replace)
            if hit:
                count += hit
                return
            full = "".join(r.text for r in para.runs)
            if find not in full:
                return
            count += full.count(find)
            new_full = full.replace(find, replace)
        else:
            pattern = _re.compile(_re.escape(find), _re.IGNORECASE)
            hit = 0
            for r in para.runs:
                ms = pattern.findall(r.text)
                if ms:
                    hit += len(ms)
                    r.text = pattern.sub(replace, r.text)
            if hit:
                count += hit
                return
            full = "".join(r.text for r in para.runs)
            if not pattern.search(full):
                return
            count += len(pattern.findall(full))
            new_full = pattern.sub(replace, full)

        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""

    for para in doc.paragraphs:
        _replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para)
    return count


def _do_docxedit(data: bytes, find: str, replace: str, case_sensitive: bool) -> tuple:
    """Orqaga moslik: bitta juftlik. Yangi kod `_do_docxedit_pairs` ishlatadi."""
    from docx import Document
    doc = Document(io.BytesIO(data))
    count = _docxedit_apply(doc, find, replace, case_sensitive)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), count


# Tashqi nom — eski test va kodlar qaytarib chaqirishi mumkin.
_docxedit_apply = _docxedit_apply_one


def _do_docxedit_pairs(data: bytes, pairs: list, case_sensitive: bool) -> tuple:
    """Bir nechta find/replace juftligini bir hujjatga ketma-ket qo'llaydi.

    pairs = [{"find": "...", "replace": "..."}, ...]
    Bo'sh `find` bo'lgan juftliklar quvib o'tkaziladi.
    Qaytaradi: (docx_bytes, jami_almashtirish_soni, per_pair_counts)
    """
    from docx import Document
    doc = Document(io.BytesIO(data))
    total = 0
    per_pair = []
    for pair in pairs:
        find = (pair.get("find") or "").strip()
        replace = pair.get("replace") or ""
        if not find:
            per_pair.append(0)
            continue
        n = _docxedit_apply_one(doc, find, replace, case_sensitive)
        per_pair.append(n)
        total += n
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), total, per_pair


_DOCXEDIT_MAX_PAIRS = 50
_DOCXEDIT_MAX_LEN = 500


def _legacy_doc_to_docx_via_libreoffice(data: bytes) -> bytes:
    """Eski .doc (1997-2003 Office formati) → .docx aylantiradi.

    python-docx faqat `.docx` (ZIP-based) ni o'qiydi. Foydalanuvchi `.doc`
    yuborsa, avval LibreOffice bilan jimgina `.docx` ga aylantiramiz, keyin
    almashtirish ishlatamiz.
    """
    import subprocess
    import shutil as _sh
    import tempfile as _tf
    from pathlib import Path

    soffice = _sh.which("soffice") or _sh.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice mavjud emas — .doc faylini avval .docx ga "
            "qo'lda saqlang yoki bizning Word→PDF xizmatidan foydalaning."
        )

    with _office_convert_sem:
        with _tf.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "input.doc")
            with open(input_path, "wb") as f:
                f.write(data)

            profile_dir = _tf.mkdtemp(prefix="lo_doctodocx_")
            try:
                env = os.environ.copy()
                env["HOME"] = profile_dir
                env.setdefault("SAL_USE_VCLPLUGIN", "svp")
                env.setdefault("LC_ALL", "C.UTF-8")
                profile_uri = Path(profile_dir).resolve().as_uri()

                cmd = [
                    soffice, "--headless", "--invisible", "--nodefault",
                    "--nologo", "--nofirststartwizard", "--norestore",
                    f"-env:UserInstallation={profile_uri}",
                    "--convert-to", "docx",
                    "--outdir", tmpdir, input_path,
                ]
                try:
                    rc, _out, _err = _run_soffice(
                        cmd, env=env, timeout=_OFFICE_CONVERT_TIMEOUT,
                    )
                except subprocess.TimeoutExpired:
                    raise RuntimeError(
                        f"LibreOffice .doc→.docx timeout {_OFFICE_CONVERT_TIMEOUT}s"
                    )

                if rc != 0:
                    stderr = (_err or b"").decode(errors="replace")[:300]
                    raise RuntimeError(
                        f"LibreOffice .doc→.docx muvaffaqiyatsiz: rc={rc}: {stderr}"
                    )

                # LibreOffice yaratgan .docx ni topamiz (nom kichik harf bilan keladi)
                docx_path = None
                for name in os.listdir(tmpdir):
                    if name.lower().endswith(".docx"):
                        docx_path = os.path.join(tmpdir, name)
                        break
                if not docx_path or not os.path.exists(docx_path):
                    raise RuntimeError("LibreOffice .docx natija chiqarmadi")
                if os.path.getsize(docx_path) < 100:
                    raise RuntimeError("LibreOffice .docx natija bo'sh")

                with open(docx_path, "rb") as f:
                    return f.read()
            finally:
                _sh.rmtree(profile_dir, ignore_errors=True)


@router.post("/api/docxedit")
@limiter.limit("15/minute")
async def docx_edit(
    request: Request,
    file: UploadFile = File(...),
    find: str = Form(""),
    replace: str = Form(""),
    case_sensitive: bool = Form(False),
    pairs: str = Form(""),
):
    """Word matnini almashtirish.

    Ikki rejim qo'llab-quvvatlanadi (orqaga moslik):
      1) Bitta juftlik (eski API): `find` + `replace` form maydonlari.
      2) Bir nechta juftlik (yangi): `pairs` JSON: [{"find": "...", "replace": "..."}].

    `pairs` mavjud bo'lsa, u ustuvor (find/replace e'tiborga olinmaydi).
    """
    import json as _json
    t0 = time.time()
    user_id = _get_user_id(request)
    try:
        data = await read_upload(file, "/api/docxedit")

        # ── Format aniqlash: .doc bo'lsa avval .docx ga aylantirish ──
        original_was_legacy_doc = False
        if _looks_like_docx(data):
            pass  # .docx — to'g'ridan-to'g'ri ishlatamiz
        elif _looks_like_legacy_doc(data):
            # python-docx eski .doc formatini o'qiy olmaydi.
            # LibreOffice orqali jimgina .docx ga aylantiramiz.
            original_was_legacy_doc = True
            try:
                loop = asyncio.get_running_loop()
                data = await asyncio.wait_for(
                    loop.run_in_executor(
                        _io_pool, _legacy_doc_to_docx_via_libreoffice, data
                    ),
                    timeout=float(_OFFICE_CONVERT_TIMEOUT + 10),
                )
            except asyncio.TimeoutError:
                raise HTTPException(
                    status_code=408,
                    detail=(
                        "Eski .doc faylni .docx ga aylantirish vaqti tugadi. "
                        "Iltimos, faylni avval Word'da .docx sifatida saqlang."
                    ),
                )
            except RuntimeError as e:
                logger.warning(f"doc2docx aylantirish xatosi: {type(e).__name__}: {str(e)[:200]}")
                raise HTTPException(
                    status_code=422,
                    detail=(
                        "Eski .doc faylni aylantirib bo'lmadi. "
                        "Iltimos, faylni Word'da .docx sifatida saqlang va qayta yuklang."
                    ),
                )
        else:
            raise HTTPException(
                status_code=415,
                detail=(
                    "Faqat .doc yoki .docx fayllar qabul qilinadi. "
                    "Yuklangan fayl Word hujjati emas."
                ),
            )

        # ── Pairs ro'yxatini tayyorlash ──
        pairs_list: list = []
        if pairs:
            try:
                parsed = _json.loads(pairs)
            except _json.JSONDecodeError:
                raise HTTPException(
                    status_code=400,
                    detail="Almashtirish ro'yxati noto'g'ri formatda yuborildi.",
                )
            if not isinstance(parsed, list):
                raise HTTPException(
                    status_code=400,
                    detail="Almashtirish ro'yxati massiv (array) bo'lishi kerak.",
                )
            for p in parsed:
                if not isinstance(p, dict):
                    continue
                f = (p.get("find") or "")
                r = (p.get("replace") or "")
                if not isinstance(f, str) or not isinstance(r, str):
                    continue
                f = f[:_DOCXEDIT_MAX_LEN]
                r = r[:_DOCXEDIT_MAX_LEN]
                if f.strip():
                    pairs_list.append({"find": f, "replace": r})
        elif find:
            if len(find) > _DOCXEDIT_MAX_LEN:
                raise HTTPException(
                    status_code=400,
                    detail=f"Qidirilayotgan matn {_DOCXEDIT_MAX_LEN} belgidan oshmasligi kerak.",
                )
            if len(replace) > _DOCXEDIT_MAX_LEN:
                raise HTTPException(
                    status_code=400,
                    detail=f"Almashtiriluvchi matn {_DOCXEDIT_MAX_LEN} belgidan oshmasligi kerak.",
                )
            pairs_list = [{"find": find, "replace": replace}]

        if not pairs_list:
            raise HTTPException(
                status_code=400,
                detail="Hech qanday qidirish matni kiritilmagan. Kamida bitta juftlik bo'lishi kerak.",
            )
        if len(pairs_list) > _DOCXEDIT_MAX_PAIRS:
            raise HTTPException(
                status_code=400,
                detail=f"Bir vaqtda maksimal {_DOCXEDIT_MAX_PAIRS} ta almashtirish ruxsat etiladi.",
            )

        loop = asyncio.get_running_loop()
        fn = functools.partial(_do_docxedit_pairs, data, pairs_list, case_sensitive)
        out_bytes, total_count, per_pair = await asyncio.wait_for(
            loop.run_in_executor(_io_pool, fn),
            timeout=45.0,
        )

        orig_name = os.path.splitext(file.filename or "document")[0]
        out_name = f"{orig_name}_edited.docx"
        enc_name = url_quote(out_name)

        if len(pairs_list) == 1:
            info_text = f"{total_count} ta so'z almashtirildi"
        else:
            info_text = (
                f"{len(pairs_list)} ta juftlik · jami {total_count} ta almashtirish"
            )
        if original_was_legacy_doc:
            info_text = f"{info_text} · .doc → .docx"

        logger.info(
            f"docxedit: {len(pairs_list)} pair(s), total={total_count}, "
            f"per_pair={per_pair}, {time.time()-t0:.1f}s"
        )
        return Response(
            content=out_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{enc_name}",
                "X-Info": safe_header(info_text),
            },
        )
    except HTTPException:
        raise
    except asyncio.TimeoutError:
        # 408 (Request Timeout) for consistency with the other convert
        # endpoints — the condition is user-fixable (smaller file / fewer pairs).
        raise HTTPException(
            status_code=408,
            detail="Almashtirish vaqti tugadi. Kichikroq fayl yoki kamroq juftlik tanlang.",
        )
    except Exception as e:
        logger.error(f"docxedit xato: {type(e).__name__}: {e}")
        raise HTTPException(
            status_code=500,
            detail="Word almashtirishda kutilmagan xato. Iltimos, qayta urinib ko'ring.",
        )


# ─── Images → PDF ─────────────────────────────────────────────────────────────

_IMG2PDF_MAX_FILES = 20


# ── Backend fallback ──────────────────────────────────────────────────────────
# /api/img2pdf and /api/imgs2pdf exist as SERVER-SIDE FALLBACK ONLY.
# The primary implementation is the client-side window.Img2PdfPro (jsPDF).
# ServicePage routes img2pdf/imgs2pdf directly to Img2PdfPro and never
# calls these endpoints under normal operation. They are kept here for
# optional server-assisted generation and testing contracts only.
# ─────────────────────────────────────────────────────────────────────────────

def _do_imgs2pdf(images_data: list) -> bytes:
    """Convert one or more images to a multi-page PDF using Pillow + pikepdf."""
    import pikepdf
    from PIL import Image, ImageOps

    pdf = pikepdf.new()
    try:
        for raw in images_data:
            # Process one image at a time and release it before the next so peak
            # RAM stays bounded (decoded pixels can be far larger than the file).
            img = Image.open(io.BytesIO(raw))
            try:
                img.load()
                try:
                    img = ImageOps.exif_transpose(img)
                except Exception:
                    pass

                if img.mode in ("RGBA", "LA"):
                    bg = Image.new("RGB", img.size, "white")
                    alpha = img.getchannel("A")
                    bg.paste(img.convert("RGBA"), mask=alpha)
                    img = bg
                elif img.mode == "P":
                    img = img.convert("RGBA")
                    bg = Image.new("RGB", img.size, "white")
                    bg.paste(img, mask=img.getchannel("A") if "A" in img.getbands() else None)
                    img = bg
                elif img.mode == "L":
                    pass
                elif img.mode != "RGB":
                    img = img.convert("RGB")

                mode = "L" if img.mode == "L" else "RGB"
                iw, ih = img.size

                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=92)
                jpeg_bytes = buf.getvalue()
            finally:
                img.close()

            colorspace = pikepdf.Name("/DeviceGray" if mode == "L" else "/DeviceRGB")
            img_obj = pikepdf.Stream(pdf, jpeg_bytes)
            img_obj.stream_dict = pikepdf.Dictionary(
                Type=pikepdf.Name("/XObject"),
                Subtype=pikepdf.Name("/Image"),
                Width=iw,
                Height=ih,
                ColorSpace=colorspace,
                BitsPerComponent=8,
                Filter=pikepdf.Name("/DCTDecode"),
            )

            page = pikepdf.Dictionary(
                Type=pikepdf.Name("/Page"),
                MediaBox=[0, 0, iw, ih],
                Resources=pikepdf.Dictionary(
                    XObject=pikepdf.Dictionary(Im0=img_obj)
                ),
                Contents=pikepdf.Stream(
                    pdf,
                    f"q {iw} 0 0 {ih} 0 0 cm /Im0 Do Q".encode(),
                ),
            )
            pdf.pages.append(pikepdf.Page(page))

        out = io.BytesIO()
        pdf.save(out, linearize=True)
        return out.getvalue()
    finally:
        pdf.close()

@router.post("/api/imgs2pdf")
@limiter.limit("5/minute")
async def imgs2pdf(request: Request, files: List[UploadFile] = File(...)):
    user_id = _get_user_id(request)
    try:
        if len(files) > _IMG2PDF_MAX_FILES:
            raise HTTPException(status_code=400, detail=f"Maksimal {_IMG2PDF_MAX_FILES} ta rasm")
        images = await read_uploads(files, "/api/imgs2pdf")
        for f, data in zip(files, images):
            if not is_image_bytes(data):
                raise HTTPException(status_code=422, detail=f"'{f.filename}' rasm fayli emas.")
        loop = asyncio.get_running_loop()
        try:
            pdf_bytes = await asyncio.wait_for(
                loop.run_in_executor(_io_pool, _do_imgs2pdf, images),
                timeout=60.0,
            )
        except asyncio.TimeoutError:
            raise HTTPException(status_code=408, detail="Rasmlar PDF ga aylantirishda vaqt tugadi.")
        logger.info(f"imgs2pdf: {len(images)} ta rasm → {len(pdf_bytes)//1024}KB PDF")
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": "attachment; filename=images.pdf",
                "X-Info": safe_header(f"{len(images)} ta rasm"),
            },
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"imgs2pdf xato: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail="Xizmatda kutilmagan xatolik yuz berdi. Iltimos, qayta urinib ko'ring.")
